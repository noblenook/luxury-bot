"""번개장터 봇 개발 사항 요약 → Word 문서 생성"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 기본 스타일 설정 ─────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2 + level * 0.25)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x20, 0x40, 0x20)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F4F0")
    p._p.get_or_add_pPr().append(shading)
    return p

def sep():
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 제목
# ════════════════════════════════════════════════════════════════
t = doc.add_heading("번개장터 자동 등록 봇 — 개발 사항 요약", level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

para("작성일: 2026-06-13   |   GitHub: https://github.com/noblenook/luxury-bot")
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
h1("1. 프로젝트 개요")
# ════════════════════════════════════════════════════════════════
para("명품 판매 플랫폼 번개장터(bunjang.co.kr)에 상품을 자동 등록하는 Python 봇입니다.")
para("product.yaml에 상품 정보를 입력하고, 이미지 폴더를 선택하면 폼 입력부터 최종 등록까지 자동으로 처리합니다.")
sep()

h2("기술 스택")
bullet("Python 3.14 + Playwright (비동기 브라우저 자동화)")
bullet("네이버 웨일 브라우저 — CDP(원격 디버깅) 방식으로 연결, 로그인 세션 유지")
bullet("YAML — 상품 정보 / 계정 정보 관리")
bullet("tkinter — 이미지 폴더 선택 GUI")
sep()

# ════════════════════════════════════════════════════════════════
h1("2. 프로젝트 파일 구조")
# ════════════════════════════════════════════════════════════════
code_block(
    "luxury-bot/\n"
    "├── main.py          # 실행 진입점\n"
    "├── bot.py           # 번개장터 봇 클래스 (BungaeBot)\n"
    "├── product.yaml     # 상품 정보 + 플랫폼별 가격\n"
    "├── credentials.yaml # 네이버 로그인 계정 (gitignore 처리)\n"
    "├── requirements.txt # 의존 패키지\n"
    "└── .gitignore       # credentials.yaml, __pycache__ 등 제외"
)
sep()

# ════════════════════════════════════════════════════════════════
h1("3. 실행 흐름")
# ════════════════════════════════════════════════════════════════

steps = [
    ("python main.py 실행", ""),
    ("tkinter 폴더 선택창", "이미지 폴더 선택"),
    ("이미지 수집", "1.jpg ~ 10.jpg 자동 수집 / product.txt 로드"),
    ("Whale 브라우저 실행", "--remote-debugging-port=9222 로 CDP 연결"),
    ("기존 탭 재사용", "browser.contexts[0].pages[0] — 빈 탭 생성 없음"),
    ("/products/new 이동", "세션 없으면 네이버 OAuth 로그인 후 재이동"),
    ("폼 자동 입력", "이미지 → 상품명 → 카테고리 → 가격 → 상품상태 → 설명"),
    ("사용자 확인 대기", '"엔터를 누르면 등록하기를 클릭합니다" — 내용 확인 가능'),
    ("등록하기 클릭", "완료 후 상품 페이지 URL 출력"),
]
for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f"  {i}.  {title}")
    r1.bold = True
    r1.font.size = Pt(10)
    if desc:
        r2 = p.add_run(f"  →  {desc}")
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
sep()

# ════════════════════════════════════════════════════════════════
h1("4. 핵심 기술 결정사항")
# ════════════════════════════════════════════════════════════════

h2("브라우저 연동")
bullet("Playwright 직접 실행 대신 Whale을 CDP(원격 디버깅)로 연결")
bullet("--user-data-dir로 Whale 실제 프로필 사용 → 쿠키/세션 자동 유지")
bullet("browser.contexts[0].pages[0]로 기존 탭 재사용 (빈 탭 생성 없음)")
bullet("wait_until=\"domcontentloaded\" 사용 — networkidle은 번개장터에서 무한 대기 발생")
sep()

h2("네이버 OAuth 로그인")
bullet("번개장터 로그인 버튼 셀렉터: button:has-text('네이버로 로그인')")
bullet("팝업 방식 → context.expect_page()로 팝업 캐치")
bullet("팝업 닫힘 이벤트(wait_for_event('close'))로 로그인 완료 감지")
bullet("2단계 인증이 있을 경우 최대 3분 수동 처리 대기")
sep()

h2("이미지 업로드")
bullet("파일 input이 hidden 속성 → wait_for() 없이 직접 set_input_files() 사용")
bullet("셀렉터: #media-input")
sep()

h2("상품 상태 드롭다운")
bullet("커스텀 드롭다운 (native <select> 아님) → select_option() 사용 불가")
bullet("'text=상품 상태를 선택해 주세요' 클릭 → 옵션 텍스트로 클릭")
sep()

h2("컨디션 등급 매핑")
tbl = doc.add_table(rows=6, cols=2)
tbl.style = "Table Grid"
headers = tbl.rows[0].cells
headers[0].text = "등급"
headers[1].text = "번개장터 표시 텍스트"
rows_data = [
    ("N", "새 상품 (미사용)"),
    ("S / A+", "사용감 없음"),
    ("A", "사용감 적음"),
    ("B", "사용감 많음"),
    ("C", "고장/파손 상품"),
]
for i, (grade, label) in enumerate(rows_data, 1):
    tbl.rows[i].cells[0].text = grade
    tbl.rows[i].cells[1].text = label
sep()

h2("상품 설명 (product.txt)")
bullet("이미지 폴더에 product.txt 함께 보관")
bullet("인코딩 자동 감지: UTF-8 → CP949 → EUC-KR 순서로 시도")
bullet("첫 줄(영문 상품명) 자동 제거 — lines[1:]")
bullet("'??? 검색 태그' 포함 줄 이후 전체 자동 제거")
sep()

# ════════════════════════════════════════════════════════════════
h1("5. 설정 파일 형식")
# ════════════════════════════════════════════════════════════════

h2("product.yaml")
code_block(
    "product:\n"
    "  title_ko: \"샤넬 클래식 플랩백 미디움 블랙 캐비어 금장\"\n"
    "  brand: \"CHANEL\"\n"
    "  condition: \"S\"   # N / S / A+ / A / B / C\n"
    "  ...\n"
    "\n"
    "prices:\n"
    "  bungae: 12300000\n"
    "  mustit: 12500000\n"
    "  # 기타 플랫폼 가격도 동일 파일에 정의"
)
sep()

h2("credentials.yaml  (gitignore — 절대 공유 금지)")
code_block(
    "bungae:\n"
    "  naver_id: \"아이디\"\n"
    "  naver_password: \"비밀번호\""
)
para("※ 메모장으로 작성 후 확장자를 .yaml로 저장하면 됩니다.", indent=True)
sep()

# ════════════════════════════════════════════════════════════════
h1("6. 주요 트러블슈팅")
# ════════════════════════════════════════════════════════════════

issues = [
    ("Whale 실행 실패",
     "executable_path 방식 → --no-startup-window로 즉시 종료됨",
     "--remote-debugging-port + connect_over_cdp 방식으로 변경"),
    ("로그인 세션 미적용",
     "new_context() 사용 시 쿠키 없음",
     "browser.contexts[0] 재사용"),
    ("빈 탭 반복 생성",
     "매 실행마다 about:blank 탭이 새로 열림",
     "context.pages[0] 재사용, about:blank 인수 제거"),
    ("networkidle 무한 대기",
     "번개장터가 networkidle 상태 미도달",
     "domcontentloaded + wait_for_timeout 조합으로 변경"),
    ("hidden 파일 input",
     "wait_for(visible) 타임아웃",
     "#media-input에 직접 set_input_files()"),
    ("UnicodeDecodeError",
     "product.txt가 EUC-KR 인코딩",
     "utf-8 → cp949 → euc-kr 순차 시도"),
    ("상품 상태 드롭다운",
     "select_option() 실패 — 커스텀 UI",
     "텍스트 클릭 방식으로 전환"),
    ("설명 첫 줄·태그 섹션 입력됨",
     "product.txt 전체가 그대로 입력됨",
     "첫 줄 제거(lines[1:]) + 검색태그 이하 제거"),
]

for title, cause, fix in issues:
    p = doc.add_paragraph()
    r = p.add_run(f"▸ {title}")
    r.bold = True
    r.font.size = Pt(10)
    bullet(f"원인: {cause}", level=1)
    bullet(f"해결: {fix}", level=1)
sep()

# ════════════════════════════════════════════════════════════════
h1("7. 다른 플랫폼 확장 시 체크리스트")
# ════════════════════════════════════════════════════════════════

para("번개장터 봇(BungaeBot)의 구조를 그대로 복사해 각 플랫폼별 파일을 생성합니다.")
para("예: guugus_bot.py, feelway_bot.py, trenbe_bot.py, rebonz_bot.py, mustit_bot.py")
sep()

h2("플랫폼별 로그인 방식")
tbl2 = doc.add_table(rows=6, cols=3)
tbl2.style = "Table Grid"
h = tbl2.rows[0].cells
h[0].text = "플랫폼"
h[1].text = "로그인 방식"
h[2].text = "비고"
platform_data = [
    ("구구스 (Guugus)", "ID / PW + 구글 2FA", "2FA는 수동 처리, 최대 120초 대기"),
    ("필웨이 (Feelway)", "ID / PW 직접", "일반 폼 로그인"),
    ("트렌비 (Trenbe)", "ID / PW 직접", "일반 폼 로그인"),
    ("리본즈 (Rebonz)", "네이버 OAuth", "번개장터와 동일 방식"),
    ("머스트잇 (Mustit)", "네이버 OAuth", "번개장터와 동일 방식"),
]
for i, (p1, p2, p3) in enumerate(platform_data, 1):
    r = tbl2.rows[i].cells
    r[0].text = p1
    r[1].text = p2
    r[2].text = p3
sep()

h2("폼 셀렉터 확인 방법")
bullet("각 플랫폼 등록 페이지에서 단계별 screenshot 찍어 셀렉터 확인")
bullet("await page.screenshot(path='debug_01.png') 코드를 각 단계에 삽입")
bullet("hidden input, 커스텀 드롭다운 여부 반드시 확인")
sep()

h2("credentials.yaml에 계정 추가")
code_block(
    "bungae:\n"
    "  naver_id: \"...\"\n"
    "  naver_password: \"...\"\n"
    "\n"
    "feelway:\n"
    "  id: \"...\"\n"
    "  password: \"...\"\n"
    "\n"
    "# 각 플랫폼별 키 이름은 main.py에서 참조하는 키와 일치해야 함"
)
sep()

# ════════════════════════════════════════════════════════════════
h1("8. GitHub 버전 관리")
# ════════════════════════════════════════════════════════════════
bullet("저장소: https://github.com/noblenook/luxury-bot")
bullet("credentials.yaml은 .gitignore에 등록 — push되지 않음")
bullet("회사 PC에서: git clone → credentials.yaml 별도 작성 → python main.py")
sep()

h2("기본 Git 워크플로우")
code_block(
    "git add .\n"
    "git commit -m \"변경 내용 설명\"\n"
    "git push\n"
    "\n"
    "# 회사 PC에서 최신 코드 받기\n"
    "git pull"
)

# ════════════════════════════════════════════════════════════════
out = r"D:\cowork\luxury-bot\번개장터봇_개발사항요약.docx"
doc.save(out)
print(f"저장 완료: {out}")
